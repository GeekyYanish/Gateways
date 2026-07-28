from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "deliverables" / "Parallax-Project-Handoff.docx"
ASSETS = ROOT / "assets"
CLIPBOARD = Path("/var/folders/9b/_xlnvdpj7cd2hxwvgdmqkw3r0000gn/T")

VIOLET = "6F3AA8"
DEEP_VIOLET = "34214F"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "27303A"
MUTED = "5E6873"
PALE_VIOLET = "F1EAF8"
PALE_BLUE = "E8EEF5"
PALE_GRAY = "F2F4F7"
WHITE = "FFFFFF"
GREEN = "2F7D5A"
GOLD = "9A6C00"
RED = "9B1C1C"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(width_dxa))
    tcW.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths, indent=120):
    """Fixed DXA table geometry, matching the business brief preset."""
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    total = sum(widths)
    tblPr = table._tbl.tblPr
    tblW = tblPr.first_child_found_in("w:tblW")
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(total))
    tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.first_child_found_in("w:tblInd")
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)
    tblInd.set(qn("w:w"), str(indent))
    tblInd.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for node in list(grid):
        grid.remove(node)
    for w in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_borders(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = tcPr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcPr.append(borders)
    for edge, attrs in kwargs.items():
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        for key, value in attrs.items():
            element.set(qn(f"w:{key}"), str(value))


def set_run_font(run, size=11, color=INK, bold=False, italic=False, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def paragraph(doc_or_cell, text="", style=None, size=None, color=INK, bold=False, italic=False,
              before=0, after=6, align=WD_ALIGN_PARAGRAPH.LEFT, keep=False):
    p = doc_or_cell.add_paragraph(style=style) if style else doc_or_cell.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    p.paragraph_format.keep_with_next = keep
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size or 11, color=color, bold=bold, italic=italic)
    return p


def clear_cell(cell):
    p = cell.paragraphs[0]
    p._element.getparent().remove(p._element)


def cell_text(cell, text, size=10.5, color=INK, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, after=0):
    clear_cell(cell)
    p = paragraph(cell, text, size=size, color=color, bold=bold, after=after, align=align)
    return p


def set_paragraph_border(p, color=VIOLET, size="12", space="4"):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_hyperlink(paragraph_obj, text, url, color=BLUE):
    part = paragraph_obj.part
    rid = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rid)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    rPr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph_obj._p.append(hyperlink)


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5 + level * 0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    r = p.add_run(text)
    set_run_font(r, 10.5)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    r = p.add_run(text)
    set_run_font(r, 10.5)
    return p


def add_image(cell_or_doc, path, width, caption=None):
    p = cell_or_doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3 if caption else 0)
    p.add_run().add_picture(str(path), width=Inches(width))
    if caption:
        c = paragraph(cell_or_doc, caption, size=8.5, color=MUTED, italic=True, after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
    return p


def add_section_title(doc, title, subtitle=None):
    p = paragraph(doc, title, size=16, color=BLUE, bold=True, before=16, after=5, keep=True)
    if subtitle:
        paragraph(doc, subtitle, size=10, color=MUTED, italic=True, after=8)
    return p


def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Parallax · Project handoff")
    set_run_font(r, 8.5, MUTED)
    p.add_run("  |  ")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    p._p.append(field)


def make_doc():
    OUT.parent.mkdir(exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    add_footer(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        s = styles[name]
        s.font.name = "Calibri"
        s._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        s._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        s.font.size = Pt(size)
        s.font.color.rgb = RGBColor.from_string(color)
        s.font.bold = True
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.line_spacing = 1.10

    # Cover — editorial cover header pattern, with a restrained Parallax violet override.
    paragraph(doc, "PARALLAX", size=10, color=VIOLET, bold=True, after=7, align=WD_ALIGN_PARAGRAPH.CENTER)
    title = paragraph(doc, "Project Design &\nImplementation Handoff", size=27, color=DEEP_VIOLET,
                      bold=True, after=5, align=WD_ALIGN_PARAGRAPH.CENTER)
    title.paragraph_format.line_spacing = 0.95
    paragraph(doc, "A concise briefing for the team: current progress, visual direction, implementation approach, and the next delivery steps.",
              size=11, color=MUTED, after=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    hero = ASSETS / "Login" / "Minecraft UI copy.png"
    add_image(doc, hero, 3.35, "Provided visual reference — landing screen / portal entry")
    status = doc.add_table(rows=1, cols=3)
    set_table_geometry(status, [3120, 3120, 3120], indent=0)
    for cell, label, value, color in [
        (status.cell(0, 0), "PROJECT STATE", "Phases 0–5 built", GREEN),
        (status.cell(0, 1), "CORE IDEA", "Voxel-styled fest portal", VIOLET),
        (status.cell(0, 2), "NEXT MILESTONE", "Supabase + organizer tools", GOLD),
    ]:
        set_cell_shading(cell, PALE_GRAY)
        set_cell_borders(cell, top={"val": "single", "sz": "4", "color": "DADCE0"}, bottom={"val": "single", "sz": "4", "color": "DADCE0"}, left={"val": "single", "sz": "4", "color": "DADCE0"}, right={"val": "single", "sz": "4", "color": "DADCE0"})
        clear_cell(cell)
        paragraph(cell, label, size=7.5, color=MUTED, bold=True, after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
        paragraph(cell, value, size=9.2, color=color, bold=True, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
    paragraph(doc, "Prepared from the current codebase and supplied design screens · 26 July 2026", size=8.5, color=MUTED, italic=True, before=12, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Page 2 — status.
    doc.add_page_break()
    add_section_title(doc, "1. Where the project is today", "The product is an interactive prototype, not just a visual concept.")
    lead = doc.add_table(rows=1, cols=1)
    set_table_geometry(lead, [9360])
    set_cell_shading(lead.cell(0, 0), PALE_VIOLET)
    set_cell_borders(lead.cell(0, 0), left={"val": "single", "sz": "14", "color": VIOLET})
    cell_text(lead.cell(0, 0), "Team takeaway: the end-to-end participant experience, the 3D/2D event world, and a repository-based data model are in place. The work now moves from a self-contained demo to production-ready content, backend security, and organizer operations.", size=11, color=DEEP_VIOLET, bold=True, after=0)
    add_section_title(doc, "Built and reviewable", None)
    for item in [
        "Portal-led visitor journey: landing → entering animation → sign-in/sign-up → character creation → travelling transition → world.",
        "Public festival pages: events, event detail, schedule, sponsors, and leaderboard; protected participant areas include dashboard, inventory, team, achievements, notifications, profile, and settings.",
        "A walkable voxel village plus equal Map and List views. Desktop-capable browsers can use 3D; phones and assistive paths retain practical alternatives.",
        "Local data system for profiles, events, registrations, teams, XP, achievements, and attendance. Its repository interface is deliberately ready to swap to Supabase.",
        "Motion and accessibility foundations: GSAP for cinematic timelines, Framer Motion for UI elements, a reduced-motion path, responsive layouts, labelled controls, and a mobile tab bar.",
    ]:
        add_bullet(doc, item)
    add_section_title(doc, "Current implementation status", None)
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [2100, 4350, 2910])
    headers = ["Area", "What exists now", "Production follow-up"]
    for i, h in enumerate(headers):
        cell_text(table.cell(0, i), h, size=9.5, color=WHITE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(table.cell(0, i), DEEP_VIOLET)
    rows = [
        ("Experience", "Portal, auth, character, world, dashboard and public routes implemented.", "Run a full team QA pass with real event copy."),
        ("World", "3D village, 2D interactive map and list access all describe the same locations.", "Finish art pass and test on target devices."),
        ("Data", "localStorage repository with idempotent XP, registration/waitlist and achievement rules.", "Replace with Supabase before trusted attendance or roles."),
        ("Art", "Manifest-based placeholders maintain the intended image dimensions and avoid layout shifts.", "Deliver original production art into public/art/**."),
    ]
    for a, b, c in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, (a, b, c)):
            cell_text(cell, value, size=9.2, color=INK)
            set_cell_borders(cell, top={"val": "single", "sz": "4", "color": "DADCE0"}, bottom={"val": "single", "sz": "4", "color": "DADCE0"}, left={"val": "single", "sz": "4", "color": "DADCE0"}, right={"val": "single", "sz": "4", "color": "DADCE0"})
    paragraph(doc, "Important scope note: localStorage is appropriate for the prototype but is not a security boundary. Do not use it for trusted QR attendance, prizes, or live organizer permissions.", size=9.2, color=RED, italic=True, before=7, after=0)

    # Page 3 — screenshot storyboard.
    doc.add_page_break()
    add_section_title(doc, "2. Intended participant experience", "The supplied screen set maps to an implementation already represented in the route structure.")
    paragraph(doc, "The design tells a clear story: enter a special realm, create an identity, then discover the festival as a world. The portal is a navigation device, not a decorative splash screen; it gives the product a memorable opening and hides route/data preparation behind a purposeful transition.", size=10.5, after=8)
    flow = doc.add_table(rows=1, cols=3)
    set_table_geometry(flow, [3120, 3120, 3120])
    images = [
        (ASSETS / "Login" / "Minecraft UI copy.png", "1. Landing — establish the Parallax premise."),
        (ASSETS / "Login" / "No_1_Minecraft UI enter animation.png", "2. Entering — portal animation provides a loading bridge."),
        (ASSETS / "Login" / "No_2_Minecraft UI welcome page.png", "3. Login / signup — a focused, game-like form."),
    ]
    for i, (img, cap) in enumerate(images):
        clear_cell(flow.cell(0, i))
        set_cell_shading(flow.cell(0, i), PALE_GRAY)
        add_image(flow.cell(0, i), img, 1.55, cap)
    flow2 = doc.add_table(rows=1, cols=3)
    set_table_geometry(flow2, [3120, 3120, 3120])
    images2 = [
        (CLIPBOARD / "codex-clipboard-f0d614fc-4301-422a-94aa-ec6335f05b84.png", "4. Create character — player name, college and avatar."),
        (ASSETS / "Login" / "No_3_Minecraft UI after login.png", "5. Travelling — transition to the festival space."),
        (ASSETS / "Homapage" / "No_4_Minecraft UI World Spawn Map.png", "6. World spawn — event categories become places to explore."),
    ]
    for i, (img, cap) in enumerate(images2):
        clear_cell(flow2.cell(0, i))
        set_cell_shading(flow2.cell(0, i), PALE_GRAY)
        width = 1.55 if i < 2 else 2.75
        add_image(flow2.cell(0, i), img, width, cap)
    add_section_title(doc, "What makes this feel coherent", None)
    for item in [
        "One stable motif: deep-violet portal energy carries the visitor from landing through authentication and into the world.",
        "The world is informational as well as atmospheric: each landmark maps to an event category, so a visitor can enter through their interest rather than a generic navigation menu.",
        "Progression is visible through avatar, XP, registrations, achievement slots and the inventory metaphor, giving repeat attendance a rewarding loop.",
    ]:
        add_bullet(doc, item)

    # Page 4 — visual system and dashboard.
    doc.add_page_break()
    add_section_title(doc, "3. Visual language and responsive UI", "Use the blocky aesthetic as a consistent interaction system, while keeping every asset and character original.")
    dashboard = doc.add_table(rows=1, cols=2)
    set_table_geometry(dashboard, [5400, 3960])
    clear_cell(dashboard.cell(0, 0))
    clear_cell(dashboard.cell(0, 1))
    set_cell_shading(dashboard.cell(0, 0), PALE_GRAY)
    set_cell_shading(dashboard.cell(0, 1), PALE_GRAY)
    add_image(dashboard.cell(0, 0), CLIPBOARD / "codex-clipboard-fe055588-f266-4e20-a2f2-902c0d66db72.png", 3.35, "Desktop inventory/dashboard reference")
    add_image(dashboard.cell(0, 1), CLIPBOARD / "codex-clipboard-753f6f6c-9864-46b9-8ce1-c8aea53b038d.png", 1.8, "Mobile quick-access reference")
    add_section_title(doc, "Design decisions to carry into every screen", None)
    visual_rows = [
        ("Surfaces", "Near-black panels, inset slots, hard borders and bevel highlights make controls feel tactile without needing copied game UI art."),
        ("Accent system", "Portal violet is the primary action/transition colour; emerald, gold, cyan and red communicate states, rarity and event categories."),
        ("Type", "Pixel display type is reserved for headings and micro-labels; readable body copy supports forms, schedules and longer event information."),
        ("Responsive pattern", "Desktop uses side navigation and multi-column cards. Mobile prioritises a tab bar, quick actions and a single-column task-focused flow."),
        ("Asset policy", "Original voxel-inspired art only. Avoid game logos, fonts, characters, skins, textures and trademark terms from other games."),
    ]
    visual = doc.add_table(rows=0, cols=2)
    set_table_geometry(visual, [2100, 7260])
    for idx, (label, detail) in enumerate(visual_rows):
        cells = visual.add_row().cells
        cell_text(cells[0], label, size=9.5, color=DEEP_VIOLET, bold=True)
        cell_text(cells[1], detail, size=9.5)
        for cell in cells:
            set_cell_borders(cell, bottom={"val": "single", "sz": "4", "color": "DADCE0"})
            if idx % 2 == 0:
                set_cell_shading(cell, PALE_BLUE)
    paragraph(doc, "Implementation reference: the Tailwind v4 tokens are centralised in src/app/globals.css, while reusable components such as BlockButton, BlockPanel, BlockInput, Hotbar, PixelAvatar and Signpost prevent style drift.", size=9.3, color=MUTED, italic=True, before=8, after=0)

    # Page 5 — implementation plan.
    doc.add_page_break()
    add_section_title(doc, "4. How the website is implemented", "The architecture is intentionally modular so visual work, interaction work and backend work can advance without a rewrite.")
    architecture = doc.add_table(rows=1, cols=3)
    set_table_geometry(architecture, [2100, 3850, 3410])
    for i, h in enumerate(["Layer", "Current implementation", "Why it matters"]):
        cell_text(architecture.cell(0, i), h, size=9.5, color=WHITE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(architecture.cell(0, i), DEEP_VIOLET)
    arch_rows = [
        ("Routes & screens", "Next.js 16 App Router separates portal, auth, public and signed-in route groups.", "The visitor journey stays understandable and routes can be protected later without redesigning screens."),
        ("UI system", "React 19 + TypeScript, Tailwind v4 design tokens, Radix dialog primitives and reusable voxel-style components.", "A single component vocabulary gives every team member a safe, repeatable way to build new pages."),
        ("Motion", "GSAP owns portal/cinematic timelines; Framer Motion handles component-level transitions. Reduced motion is supported in code and settings.", "Motion adds atmosphere while remaining optional and accessible."),
        ("World", "React Three Fiber/three.js renders a generated walkable village; Map and List expose the same event locations.", "3D is a differentiator, but no visitor is blocked by device capability or motion preference."),
        ("Data seam", "All screens use an async Repository interface; LocalRepository currently stores data in localStorage.", "Supabase can replace the implementation without rewriting every screen call."),
        ("Art delivery", "A central asset manifest declares dimensions and serves generated placeholders until files arrive.", "Artists can deliver incrementally without layout shift or hard-coded paths."),
    ]
    for layer, current, why in arch_rows:
        cells = architecture.add_row().cells
        for cell, text in zip(cells, (layer, current, why)):
            cell_text(cell, text, size=8.9)
            set_cell_borders(cell, top={"val": "single", "sz": "4", "color": "DADCE0"}, bottom={"val": "single", "sz": "4", "color": "DADCE0"}, left={"val": "single", "sz": "4", "color": "DADCE0"}, right={"val": "single", "sz": "4", "color": "DADCE0"})
        set_cell_shading(cells[0], PALE_BLUE)
        cell_text(cells[0], layer, size=8.9, color=DEEP_VIOLET, bold=True)
    add_section_title(doc, "Practical development approach", None)
    for item in [
        "Build each new feature as a route-level screen composed from the shared voxel components; avoid one-off CSS that bypasses the theme tokens.",
        "Add or refine an event category in the data model, then map it to a world location and a biome scene so public listings, map labels and world markers stay aligned.",
        "Supply art to the documented public/art/** paths at the declared pixel dimensions. The placeholder system ensures screens do not jump when each asset is replaced.",
        "Before trusted live operations, connect Supabase Auth/Postgres, move route protection server-side, and enforce roles through database RLS—not through client-side UI checks.",
    ]:
        add_number(doc, item)

    # Page 6 — team plan and sources.
    doc.add_page_break()
    add_section_title(doc, "5. Recommended next steps", "Sequence the delivery so the team can demo confidently now and harden for a real event later.")
    nexts = doc.add_table(rows=1, cols=3)
    set_table_geometry(nexts, [1600, 4900, 2860])
    for i, h in enumerate(["Priority", "Work to schedule", "Definition of done"]):
        cell_text(nexts.cell(0, i), h, size=9.5, color=WHITE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(nexts.cell(0, i), DEEP_VIOLET)
    steps = [
        ("Now", "Content and experience QA: enter real event details, verify the journey on desktop/mobile, and decide the first public event categories.", "A shareable internal demo with approved copy and tested flows."),
        ("Next", "Original art production: portal scenes, five character archetypes, biome layers, sprites, badges and world-map image to the supplied manifest sizes.", "Placeholders are replaced incrementally with no visual re-layout."),
        ("Before live", "Supabase migration: auth, Postgres tables, RLS roles, server route protection and environment configuration.", "Cross-device accounts and secure data mutations."),
        ("Then", "Operational tools: full event CRUD, teams, organizer/admin dashboards, realtime announcements and QR check-in with server-held signing.", "A festival operations workflow that can be trusted on event day."),
        ("Launch hardening", "Accessibility, performance, device testing, load testing, monitoring and a final design consistency review.", "A reliable public launch, not just a polished demo."),
    ]
    for pri, work, done in steps:
        cells = nexts.add_row().cells
        for cell, text in zip(cells, (pri, work, done)):
            cell_text(cell, text, size=9.1)
            set_cell_borders(cell, top={"val": "single", "sz": "4", "color": "DADCE0"}, bottom={"val": "single", "sz": "4", "color": "DADCE0"}, left={"val": "single", "sz": "4", "color": "DADCE0"}, right={"val": "single", "sz": "4", "color": "DADCE0"})
        set_cell_shading(cells[0], PALE_BLUE)
        cell_text(cells[0], pri, size=9.1, color=DEEP_VIOLET, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_section_title(doc, "Team reference links and source material", "The following repository documents are the source of truth for implementation decisions.")
    sources = [
        ("Project overview", ROOT / "README.md", "Routes, tech stack, shipped phases, verification notes and roadmap."),
        ("3D world", ROOT / "VOXEL-3D.md", "World controls, device fallbacks, performance constraints and extension points."),
        ("Art contract", ROOT / "ART-ASSETS.md", "Original-art rules, exact output dimensions, file paths and scene briefs."),
        ("Backend plan", ROOT / "SUPABASE-MIGRATION.md", "Auth, database schema, roles/RLS and the pre-live migration plan."),
        ("Motion rules", ROOT / "ANIMATION.md", "Ownership of timelines, reduced-motion behaviour and known animation pitfalls."),
    ]
    src_table = doc.add_table(rows=0, cols=2)
    set_table_geometry(src_table, [2600, 6760])
    for title_text, file_path, detail in sources:
        cells = src_table.add_row().cells
        clear_cell(cells[0])
        p = paragraph(cells[0], after=0)
        add_hyperlink(p, title_text, file_path.as_uri(), color=BLUE)
        cell_text(cells[1], detail, size=9.2)
        for cell in cells:
            set_cell_borders(cell, bottom={"val": "single", "sz": "4", "color": "DADCE0"})
    paragraph(doc, "Visual references used in this handoff are the supplied Parallax screen images in assets/Login, assets/Homapage and the two supplied dashboard/mobile UI references. They are direction references; the product must continue using original assets and naming.", size=9, color=MUTED, italic=True, before=9, after=0)

    doc.core_properties.title = "Parallax — Project Design & Implementation Handoff"
    doc.core_properties.subject = "Current progress, visual direction and implementation approach"
    doc.core_properties.author = "Parallax project team"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    make_doc()
